from docx import Document
from docx.shared import Pt, Inches
from num2words import num2words

document = Document()


def add_centered_paragraph(doc, text):
    # paragraph = doc.add_paragraph(text)
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    style = paragraph.style
    font = style.font
    font.size = Pt(10)
    # paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.left_indent = Inches(2)


def add_title(doc, text):
    # paragraph = doc.add_paragraph(text)
    paragraph = doc.add_paragraph()
    paragraph.add_run(text).bold = True

    paragraph.paragraph_format.left_indent = Inches(2)


def create_table():
    table = document.add_table(rows=1, cols=7)
    # определяем стиль таблицы
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Наименование"
    hdr_cells[2].text = "Код"
    hdr_cells[3].text = "Ед. изм."
    hdr_cells[4].text = "Кол-во"
    hdr_cells[5].text = "Цена"
    hdr_cells[6].text = "Стоимость"

    for number, name, code, eizm, count, price, cost in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(number)
        row_cells[1].text = name
        row_cells[2].text = code
        row_cells[3].text = eizm
        row_cells[4].text = count
        row_cells[5].text = price
        row_cells[6].text = cost


add_centered_paragraph(document, "Корма Лукаши, ИНН:")
add_centered_paragraph(document, "Адрес: ЛО, Гатч. рн. Пудомягское СП, Борский массив, зем. уч №1:")
add_centered_paragraph(document, "Телефон: +7(921)953-5499")

document.add_heading("", 2)
add_title(document, "Товарный чек №977 от 02.03.2024 14:20:45")

records = (
    (3, '101', 'Spam', '20', '20', '20', '20'),
    (7, '422', 'Eggs', '20', '20', '20', '20'),
    (4, '631', 'Spam, spam, eggs, and spam', '20', '20', '20', '20')
)

create_table()

document.add_paragraph("Всего наименований: 3, на сумму 2 850,00р")
document.add_paragraph("(" + num2words(2558, lang='ru') + " 00 копеек)")
document.add_paragraph("Отпустил: _______________________________")

document.save("productReceipt.docx")
