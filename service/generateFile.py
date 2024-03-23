import os
from datetime import datetime
import random
from buttons.users.main import main
from config import bot
from docx import Document
from docx.shared import Pt, Inches
from num2words import num2words


async def create_product_receipt(state, user_id):
    document = Document()

    async def add_centered_paragraph(doc, text):
        paragraph = doc.add_paragraph()
        paragraph.add_run(text)
        style = paragraph.style
        font = style.font
        font.size = Pt(10)
        paragraph.paragraph_format.left_indent = Inches(2)

    async def add_title(doc, text):
        paragraph = doc.add_paragraph()
        paragraph.add_run(text).bold = True
        paragraph.paragraph_format.left_indent = Inches(2)

    async def create_table():
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "№"
        hdr_cells[1].text = "Наименование"
        hdr_cells[3].text = "Ед. изм."
        hdr_cells[4].text = "Кол-во"
        hdr_cells[5].text = "Цена"
        hdr_cells[6].text = "Стоимость"

        for number, name, code, eizm, count, price, cost in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(number)
            row_cells[1].text = name
            row_cells[2].text = code
            row_cells[3].text = eizm
            row_cells[4].text = count
            row_cells[5].text = price
            row_cells[6].text = cost

    await add_centered_paragraph(document, "Корма Лукаши, ИНН:")
    await add_centered_paragraph(document, "Адрес: ЛО, Гатч. рн. Пудомягское СП, Борский массив, зем. уч №1:")
    await add_centered_paragraph(document, "Телефон: +7(921)953-5499")

    document.add_heading("", 2)
    random_count = random.randint(1, 10000)
    current_date = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    await add_title(document, f"Товарный чек №{random_count} от {current_date}")

    async with state.proxy() as data:
        total_price = 0
        for _, price, count_in_cart, _ in data['products'].values():
            if count_in_cart > 0:
                tp = count_in_cart * price
                total_price += tp
                records = []
                row_number = 1
                for title, price, count_in_cart, info in data['products'].values():
                    records.append((
                        row_number,
                        title,
                        info,
                        f'кг/шт',
                        f'{count_in_cart}',
                        f'{price}',
                        f'{count_in_cart * price} руб.'
                    ))
                    row_number += 1

        await create_table()

        if total_price > 5000:
            discount = total_price * 0.02
            total_price -= discount
            document.add_paragraph(f"Сумма со скидкой: {total_price} руб. (скидка 2%)")
        else:
            document.add_paragraph(f"Сумма без скидки: {total_price} руб.")

        document.add_paragraph(f"Всего наименований: {row_number - 1}, на сумму {total_price}р")
        document.add_paragraph("(" + num2words(f"{int(total_price)}", lang='ru') + ")")
        document.add_paragraph("Отпустил: _______________________________")

    async with state.proxy() as data:
        document.add_paragraph(f"Получатель: {data['name']}")
        document.add_paragraph(f"Способ оплаты: {data['pay']}\n")
        document.add_paragraph(f"Цвет машины: {data['color_auto']}")
        document.add_paragraph(f"Марка машины: {data['brand_auto']}")
        document.add_paragraph(f"Номер машины: {data['number_auto']}")

    document.save("productReceipt.docx")
    doc = f"{os.getcwd()}/productReceipt.docx"
    doc_path = open(doc, 'rb+')
    await bot.send_document(user_id, document=doc_path, reply_markup=main())
    return f"{os.getcwd()}/productReceipt.docx"



